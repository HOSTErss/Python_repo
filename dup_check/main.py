import os

import docx

base = r'D:\Code-source\python_source\dup_check\test_data'  # docx文件所在目录

doc_file_name = os.listdir(base)  # 获取docx文件名称
each_file = []
for ele in doc_file_name:
    abs_path = os.path.join(base, ele)
    each_file.append(abs_path)

length = len(each_file)
for i in range(length-1):  # 待比较的文件
    sim_list = []
    file1 = docx.Document(each_file[i])
    intent_list1 = []
    for fp in file1.paragraphs:
        # if count == 7:
        #     break
        tmp = fp.text.replace(' ', '')  # 丢弃无效内容
        if len(tmp) > 0:
            intent_list1.append(tmp)
    len1 = len(intent_list1)
    for j in range(i+1, length):  # 遍历进行比较的文件
        file2 = docx.Document(each_file[j])
        intent_list2 = []
        for fp in file2.paragraphs:
            # if count == 7:
            #     break
            tmp = fp.text.replace(' ', '')
            if len(tmp) > 0:
                intent_list2.append(tmp)

        len2 = len(intent_list2)
        # print(len1)
        # print(len2)
        sum_div = 0
        for k in range(len1):
            divisor = 0.0
            s1 = intent_list1[k]
            set1 = set(s1)
            len11 = len(set1)
            if k-5 >= 0:
                begin = k-5
            else:
                begin = 0
            if k + 5 <= len2:
                end = k + 5
            else:
                end = len2

            for s in range(begin, end):
                # print(s1)
                s2 = intent_list2[s]
                set2 = set(s2)

                set3 = set1 & set2
                len13 = len(set3)

                divisor = max(len13 / len11, divisor)
                if divisor >= 0.99:
                    # print(1)
                    break
            sum_div += divisor
        if sum_div / len1 > 0.85:
            t1 = (j, sum_div / len1)
            sim_list.append(t1)
    if len(sim_list) != 0:
        print("{}: ".format(doc_file_name[i]), end="")
        for ele in sim_list:
            print('[{}, {:.2f}];'.format(doc_file_name[ele[0]], ele[1]), end=' ')
        print()
        # print("{}-{}, {:.2f}".format(i, j, sum_div / len1))
